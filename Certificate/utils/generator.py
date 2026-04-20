import os
import time
import pandas as pd
from tqdm import tqdm
from docx import Document
from datetime import datetime
from django.conf import settings
from django.core.files.storage import default_storage
from Training.models import TrainingProgram
from Certificate.models import Certificate
from Login.models import User
from Certificate.utils.utils import replace_placeholders
from zipfile import ZipFile


# ✅ Safe flag for PDF conversion (disabled in Linux/Render)
WINDOWS_COM_AVAILABLE = False


def convert_docx_to_pdf(docx_path, pdf_path):
    print("⚠️ PDF conversion disabled (Linux environment)")
    return False


def generate_certificates_from_excel(file_path, template_path, training_code, coordinator_user):
    """
    Linux-safe version (PDF conversion disabled)
    """

    # 🚫 Skip entire feature if COM not available
    if not WINDOWS_COM_AVAILABLE:
        print("⚠️ Certificate generation skipped (not supported on this environment)")
        return [], None

    try:
        df = pd.read_excel(file_path)
    except Exception as e:
        print(f"❌ Error reading Excel: {e}")
        return [], None

    try:
        training = TrainingProgram.objects.get(code=training_code)
    except TrainingProgram.DoesNotExist:
        print(f"❌ Training with code '{training_code}' not found.")
        return [], None

    if not os.path.exists(template_path):
        print(f"❌ Template not found at: {template_path}")
        return [], None

    output_dir = os.path.join(settings.MEDIA_ROOT, 'certificates')
    os.makedirs(output_dir, exist_ok=True)

    generated_certificates = []

    for index, row in tqdm(df.iterrows(), total=len(df)):

        ehrms_code = row.get('ehrms_code') or row.get('EHRMS')
        email = row.get('email') or row.get('Email')

        user = None
        if ehrms_code:
            user = User.objects.filter(ehrms_code=ehrms_code).first()
        if not user and email:
            user = User.objects.filter(email=email).first()

        if not user:
            print(f"⚠️ Skipping row — no matching user for EHRMS: {ehrms_code}, Email: {email}")
            continue

        try:
            doc = Document(template_path)

            name = f"{user.first_name} {user.middle_name or ''} {user.last_name}".strip()
            designation = user.designation or ''
            branch = user.branch or ''
            institution = user.institute_name or ''
            year = training.start_date.year if training.start_date else datetime.now().year
            start_date = training.start_date or datetime.now()

            day = start_date.strftime("%d")
            month = start_date.strftime("%m")
            year_suffix = start_date.strftime("%y")
            serial = str(index + 1).zfill(2)

            reference_number = f"{training.code}-{day}{month}{year_suffix}-{serial}"

            replacements = {
                '{{name of staff}}': name,
                '{{designation}}': designation,
                '{{branch}}': branch,
                '{{institute name}}': institution,
                '{{certificate no}}': reference_number,
            }

            for p in doc.paragraphs:
                replace_placeholders(p, replacements)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_placeholders(p, replacements)

            docx_name = f"{user.ehrms_code}_{training.code}_{year}.docx"
            docx_path = os.path.join(output_dir, docx_name)

            doc.save(docx_path)

            print(f"⚠️ Skipping PDF conversion for {user.email}")

            # Optional: create DB entry without file
            cert = Certificate.objects.create(
                user=user,
                training=training,
                uploaded_by=coordinator_user,
                full_name=name,
                designation=designation,
                institution=institution,
                reference_number=reference_number,
            )

            from Enrollment.models import Enrollment
            enrollment, created = Enrollment.objects.get_or_create(
                trainee=user,
                training=training,
                defaults={'status': 'attended'}
            )
            if not created and enrollment.status != 'attended':
                enrollment.status = 'attended'
                enrollment.save()

            generated_certificates.append(cert)

        except Exception as e:
            print(f"❌ Error for user {user.email or ehrms_code}: {e}")
            continue

    zip_file_path = create_zip_for_training(training_code)

    return generated_certificates, zip_file_path


def create_zip_for_training(training_code):
    cert_folder = os.path.join(settings.MEDIA_ROOT, "certificates")
    zip_folder = os.path.join(settings.MEDIA_ROOT, "certificates", "zips")
    os.makedirs(zip_folder, exist_ok=True)

    zip_file_path = os.path.join(zip_folder, f"{training_code}_certificates.zip")

    with ZipFile(zip_file_path, 'w') as zipf:
        for filename in os.listdir(cert_folder):
            if filename.startswith(training_code) and filename.endswith('.pdf'):
                filepath = os.path.join(cert_folder, filename)
                zipf.write(filepath, arcname=filename)

    return zip_file_path
